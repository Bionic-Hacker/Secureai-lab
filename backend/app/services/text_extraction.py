"""
Text extraction for the RAG ingestion pipeline.

Runs entirely on the decrypted plaintext already in memory from
document_service.get_document_content — nothing here touches disk, and
nothing here re-derives trust: by the time extract_text() is called, the
document has already passed malware scanning and content-type validation
in Phase 2. This module's only job is turning bytes into a string.
"""
import io

from app.services.document_service import DocumentError


def extract_text(data: bytes, ext: str) -> str:
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    # .txt and every source-code extension (.py, .js, .ts, .java, .go,
    # .rb, .php) are all just plaintext — decode directly.
    return _extract_plaintext(data)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentError("Could not extract text from this PDF — it may be corrupted or image-only.", 422) from exc

    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        # Common for scanned/image-only PDFs with no embedded text layer.
        # OCR is a real future enhancement, not something silently faked here.
        raise DocumentError("No extractable text found in this PDF (it may be a scanned image without OCR).", 422)
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise DocumentError("Could not read this DOCX file — it may be corrupted.", 422) from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables carry real content in a lot of real-world documents (specs,
    # comparison sheets) — skipping them would silently drop information
    # a search over this document should have been able to find.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)
    if not text.strip():
        raise DocumentError("No extractable text found in this document.", 422)
    return text


def _extract_plaintext(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentError("File is not valid UTF-8 text.", 422) from exc
